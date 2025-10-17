#!/usr/bin/env python3
# coding=utf-8

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import time
import json
from decimal import Decimal
from stat_types import *

class DocxGenerator(object):
    def __init__(self, bench_eval:BenchmarkResult, file_path:str) -> None:
        self.__bench_eval = bench_eval
        self.__file_path = file_path
        self.__doc = Document()

    def __generate_header(self):
        self.__doc.add_heading('RSCL Performance Benchmark', 0)
        # insert table of content
        # need to right-click to update manually
        paragraph = self.__doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar') # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin') # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve') # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # change 1-3 depending on heading levels you need
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        self.__doc.add_page_break()

    def __generate_env(self):
        self.__doc.add_heading('一、测试环境', level=1)
        p = self.__doc.add_paragraph()
        p.add_run(self.__bench_eval.environment).bold = True

    def __generate_case_status(self):
        self.__doc.add_heading('二、测试概览', level=1)
        if len(self.__bench_eval.cases) > 0:

            table = self.__doc.add_table(rows=1, cols=3, style = 'Table Grid')
            table.autofit = False

            col_width_dict = {0:1, 1:1, 2:4.5}
            for col_num in range(3):
                table.columns[col_num].width = Inches(col_width_dict[col_num])

            is_total_passed = True
            for case in self.__bench_eval.cases:
                for _, pub_stat in case.pub_stats.items():
                    if pub_stat.exit_code != "0":
                        is_total_passed = False
                        break
                for _, sub_stat in case.sub_stats.items():
                    if sub_stat.exit_code != "0":
                        is_total_passed = False
                        break
            row_cells = table.rows[0].cells
            row_cells[0].text = 'case名称'
            row_cells[1].text = "测试结果"
            row_cells[2].text = "传输性能"
            for case in self.__bench_eval.cases:
                row_cells = table.add_row().cells
                row_cells[0].text = case.case_name
                sub_table = row_cells[2].add_table(rows=10, cols=len(case.sub_stats) + 1)
                sub_count = 0
                ret_str = ""

                for ppid, pub_stat in case.pub_stats.items():
                    if pub_stat.exit_code != "0":
                        signal = pub_stat.exit_code
                        if pub_stat.exit_code in SIGNAL_TABLE:
                            signal = SIGNAL_TABLE[pub_stat.exit_code]
                        ret_str += "\n{} Pub recv signal:".format(ppid) + signal

                for spid, sub_stat in case.sub_stats.items():
                    if sub_stat.transport_stat is None:
                        print(
                            "case {} sub pid {} has no transport_summary log, maybe unexpected exit? skip"
                            .format(case.case_name, spid))
                        continue
                    signal = sub_stat.exit_code
                    if sub_stat.exit_code in SIGNAL_TABLE:
                        signal = SIGNAL_TABLE[sub_stat.exit_code]
                    if sub_stat.exit_code != "0":
                        ret_str += "\n{} Sub recv signal:".format(spid) + signal
                    if float(sub_stat.transport_stat['平均发送耗时(ms)']) > 10:
                        ret_str += "\n{} Publish cost ms exceed 10".format(spid)
                    if float(sub_stat.transport_stat['平均接收延迟(ms)']) > 10:
                        ret_str += "\n{} Latency ms exceed 10".format(spid)
                    if len(sub_stat.transport_stat["丢帧"]) > 0:
                        ret_str += "\n{} Transport drops {} msg".format(spid, len(sub_stat.transport_stat["丢帧"]))
                    if len(sub_stat.transport_stat["乱序"]) > 0:
                        ret_str += "\n{} Transport disorder {} msg".format(spid, len(sub_stat.transport_stat["乱序"]))

                    sub_count += 1
                    key_row = 1
                    sub_table.columns[0].cells[0].text = "pid"
                    sub_table.columns[sub_count].cells[0].text = spid
                    for k, v in sub_stat.transport_stat.items():
                        value = str(v)
                        if k == u"最后一帧":
                            seq_list = re.findall(r"'([^']*)'", str(v))
                            seq_list = [int(seq.rsplit(":", 1)[-1]) for seq in seq_list]
                            value=str(max(seq_list, default=0))
                        if k == u"丢帧" or k == u"乱序":
                            value = str(len(v))
                            # example: ["cn0614007308l_31569:1699326556091082730:196"]
                            seq_list = re.findall(r"'([^']*)'", str(v))
                            seq_list = [int(seq.rsplit(":", 1)[-1]) for seq in seq_list]
                            if len(seq_list) != 0:
                                value += " (ranges from {} to {})".format(min(seq_list), max(seq_list))
                        key_cells = sub_table.columns[0].cells
                        col_cells = sub_table.columns[sub_count].cells

                        key_cells[key_row].text = k
                        col_cells[key_row].text = value
                        key_row += 1

                if len(ret_str) == 0:
                    ret_str = "Pass"
                row_cells[1].text = ret_str
        else:
            self.__doc.add_paragraph('没有发现测试结论！')
        p = self.__doc.add_paragraph()
        self.__write_image_line(p, "详情参考transport_summary_${node_topic}.${pid}.log")
        idx_map = {}
        idx_map["publish"] = 0
        idx_map["latency"] = 0
        for tune_file in self.__bench_eval.tune_files:
            if "avg_ms" in tune_file:
                pic = self.__doc.add_picture(
                        tune_file,
                        width=Inches(5.8))
                p = self.__doc.add_paragraph()
                if "publish" in tune_file:
                    idx_map["publish"] += 1
                    self.__write_image_line(p, "* publish耗时随频率和消息大小变化3d曲面图视角{}".format(idx_map["publish"]))
                if "latency" in tune_file:
                    idx_map["latency"] += 1
                    self.__write_image_line(p, "* 接收延迟随频率和消息大小变化3d曲面图视角{}".format(idx_map["latency"]))
                self.__doc.add_paragraph()
        self.__doc.add_page_break()

    def __generate_case_doc(self, idx:int, single_case:BenchmarkCaseResult):
        self.__doc.add_heading('3.{}、{}'.format(str(idx), single_case.case_name), level=2)
        case_description = ""
        for k, v in single_case.sub_stats.items():
            if len(v.description) > 0:
                case_description = v.description
                break
        if len(case_description) > 0:
            p = self.__doc.add_paragraph()
            p.add_run(case_description)
        self.__doc.add_heading('3.{}.0、Case运行时整体系统资源统计'.format(str(idx)), level=3)
        for _, sub_stat in single_case.sub_stats.items():
            if os.path.exists(sub_stat.sys_resource.plot_path):
                self.__doc.add_picture(
                    sub_stat.sys_resource.plot_path,
                    width=Inches(5.8))
                p = self.__doc.add_paragraph()
                self.__write_image_line(
                    p, '* 数据说明：统计case运行时系统整体cpu占用情况\n')
                self.__write_image_line(
                    p, '* file:{}'.format(sub_stat.sys_resource.csv_path))
                break
        self.__doc.add_heading('3.{}.1、通信性能统计'.format(str(idx)), level=3)
        p = self.__doc.add_paragraph()
        p.add_run("通信性能统计").bold = True
        for spid, sub_stat in single_case.sub_stats.items():
            p = self.__doc.add_paragraph("************* node_topic:{} *************".format(sub_stat.node_topic))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(sub_stat.transport_eval.plot_path):
                self.__doc.add_picture(
                    sub_stat.transport_eval.plot_path,
                    width=Inches(5.8))
                p = self.__doc.add_paragraph()
                self.__write_image_line(
                    p, '* 数据说明：统计通信过程中进程{}的传输性能指标。\
                    \npublish:执行Publish函数的耗时，主要是序列化。\nlatency: 接收延迟，从序列化发送到收到数据反序列化后的耗时。\
                    \njitter: 抖动，反应数据包传输的时延变化情况，抖动越低表示通信稳定性越高\n'.format(spid))
                self.__write_image_line(
                    p, '* file:{}'.format(sub_stat.transport_eval.csv_path))
        self.__doc.add_heading('3.{}.2、Publish性能统计'.format(str(idx)), level=3)
        for ppid, pub_stat in single_case.pub_stats.items():
            self.__doc.add_heading('3.{}.2.1、{} 测试配置'.format(str(idx), ppid), level=4)
            p = self.__doc.add_paragraph()
            p.add_run(json.dumps(pub_stat.test_setup, indent=4, ensure_ascii=False))
            self.__doc.add_heading('3.{}.2.2、{} 进程资源占用'.format(str(idx), ppid), level=4)
            if os.path.exists(pub_stat.proc_resource.plot_path):
                self.__doc.add_picture(
                    pub_stat.proc_resource.plot_path,
                    width=Inches(5.8))
                p = self.__doc.add_paragraph()
                self.__write_image_line(
                    p, '* 数据说明：发送数据进程{}的资源占用，包括cpu，进程数，实际使用物理内存\n'.format(ppid))
                self.__write_image_line(
                    p, '* file:{}'.format(pub_stat.proc_resource.csv_path))
        self.__doc.add_heading('3.{}.3、Subscribe性能统计'.format(str(idx)), level=3)
        for spid, sub_stat in single_case.sub_stats.items():
            self.__doc.add_heading('3.{}.3.1、{} 测试配置'.format(str(idx), spid), level=4)
            p = self.__doc.add_paragraph()
            p.add_run(json.dumps(sub_stat.test_setup, indent=4, ensure_ascii=False))
            self.__doc.add_heading('3.{}.3.2、{} 进程资源占用'.format(str(idx), spid), level=4)
            if os.path.exists(sub_stat.proc_resource.plot_path):
                self.__doc.add_picture(
                    sub_stat.proc_resource.plot_path,
                    width=Inches(5.8))
                p = self.__doc.add_paragraph()
                self.__write_image_line(
                    p, '* 数据说明：接收数据进程{}的资源占用，包括cpu，进程数，实际使用物理内存\n'.format(spid))
                self.__write_image_line(
                    p, '* file:{}'.format(sub_stat.proc_resource.csv_path))


    def generate_docx(self) -> None:
        self.__generate_header()
        self.__generate_env()
        self.__generate_case_status()

        self.__doc.add_heading(
            '三、case详情',
            level=1).style.paragraph_format.line_spacing_rule = \
                WD_LINE_SPACING.AT_LEAST
        for idx, case in enumerate(self.__bench_eval.cases, start=1):
            self.__generate_case_doc(idx, case)
        self.__doc.save(self.__file_path)

    def __write_image_line(self, paragraph, text) -> None:
        run = paragraph.add_run(text)
        run.italic = True
        run.font.size = Pt(6)
        paragraph.paragraph_format.space_before=Pt(2)
        paragraph.paragraph_format.space_after=Pt(2)

    def __write_file_link(self, document, text) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after=Pt(2)
        paragraph.paragraph_format.space_before=Pt(2)

    def __add_res_table(self, document, category, colume_title, flag=True) -> None:
        table = document.add_table(rows=1, cols=len(colume_title)*2-1)
        row_cells = table.rows[0].cells
        for i in range(len(colume_title)-1):
            row_cells[2*i+1].text = colume_title[i+1]
            row_cells[2*i+2].text = ('超过' if flag else '低于') + '阈值的时间段'
        if not self.__result.__contains__(category):
            return
        for proc, data in self.__result[category].items():
            row_cells = table.add_row().cells
            row_cells[0].text = proc
            i = 1
            for index in data.values():
                row_cells[i].text = str(index['value'])
                text = ''
                for interval in index['data']:
                    text += '{}-{}, '.format(
                        datetime.datetime.fromtimestamp(
                            interval[0]).strftime('%H:%M:%S'),
                        datetime.datetime.fromtimestamp(
                            interval[1]).strftime('%H:%M:%S'))
                row_cells[i+1].text = text
                i += 2

if __name__ == '__main__':
    pass
